import os
import subprocess
import io
import re
import math
import tempfile
import uuid
from shutil import which
import streamlit as st
import streamlit.components.v1 as components
import speech_recognition as sr
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from groq import Groq
import google.generativeai as genai
from firebase_admin import firestore
from database import db, get_user, get_system_config, increment_api_usage, get_active_keys, invalidate_user_cache

# ==========================================
# 2. FUNGSI PENDUKUNG (DOCX, FFMPEG)
# ==========================================
project_folder = os.getcwd()
local_ffmpeg, local_ffprobe = os.path.join(project_folder, "ffmpeg.exe"), os.path.join(project_folder, "ffprobe.exe")

# PENYESUAIAN KHUSUS RAILWAY / LINUX (Mencari di lokal Windows atau sistem Linux)
if os.path.exists(local_ffmpeg) and os.path.exists(local_ffprobe):
    ffmpeg_cmd, ffprobe_cmd = local_ffmpeg, local_ffprobe
    os.environ["PATH"] += os.pathsep + project_folder
else:
    # Railway biasanya menginstal di /usr/bin/ atau bisa dideteksi via which
    ffmpeg_cmd = which("ffmpeg") or "/usr/bin/ffmpeg"
    ffprobe_cmd = which("ffprobe") or "/usr/bin/ffprobe"

# Verifikasi Terakhir agar aplikasi tidak crash saat proses transkrip
if not os.path.exists(ffmpeg_cmd) and not which("ffmpeg"):
    st.error("❌ FFmpeg not found. Pastikan NIXPACKS_APT_PKGS di Railway sudah diset ke 'ffmpeg'.")
    st.stop()

def get_duration(file_path):
    try: return float(subprocess.check_output([ffprobe_cmd, "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", file_path], stderr=subprocess.STDOUT))
    except: return 0.0

def create_docx(text, title):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    import io
    
    doc = Document()
    doc.add_heading(title, level=1)
    
    in_table = False
    table_obj = None
    
    for line in text.split('\n'):
        line_str = line.strip()
        
        # Abaikan baris yang kosong
        if not line_str: 
            in_table = False
            continue
            
        # --- FITUR DETEKSI & RENDER TABEL (KHUSUS RTL & QNA) ---
        if line_str.startswith('|') and line_str.endswith('|'):
            cells = [c.strip() for c in line_str.strip('|').split('|')]
            
            if len(cells) > 0 and all(re.match(r'^[-:\s]+$', c) for c in cells):
                continue 
                
            if not in_table:
                in_table = True
                table_obj = doc.add_table(rows=1, cols=len(cells))
                table_obj.style = 'Table Grid' 
                
                hdr_cells = table_obj.rows[0].cells
                for i, val in enumerate(cells):
                    if i < len(hdr_cells):
                        clean_val = val.replace('**', '').replace('*', '')
                        hdr_cells[i].text = clean_val
                        if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
                            hdr_cells[i].paragraphs[0].runs[0].bold = True
            else:
                row_cells = table_obj.add_row().cells
                for i, val in enumerate(cells):
                    if i < len(row_cells):
                        clean_val = val.replace('**', '').replace('*', '')
                        row_cells[i].text = clean_val
            continue 
        else:
            in_table = False 
        
        # --- PARSING TEKS NORMAL ---
        if re.match(r'^\s*---\s*$', line):
            doc.add_paragraph("_" * 50)
            continue
        
        heading_match = re.match(r'^(#+)\s+(.*)', line_str)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=min(level, 9))
            continue
            
        bullet_match = re.match(r'^(\s*)[\*\-\+]\s+(.*)', line)
        number_match = re.match(r'^(\s*)([A-Za-z0-9]+[\.\)])\s+(.*)', line)
        
        p = None
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            try:
                style_name = 'List Bullet 2' if indent_spaces >= 2 else 'List Bullet'
                p = doc.add_paragraph(style=style_name)
            except:
                p = doc.add_paragraph(style='List Bullet')
            line_content = bullet_match.group(2)
            
        elif number_match:
            indent_spaces = len(number_match.group(1))
            p = doc.add_paragraph()
            if indent_spaces > 0:
                try:
                    p.paragraph_format.left_indent = Pt(18) 
                except: pass
            line_content = number_match.group(2) + " " + number_match.group(3)
            
        else:
            p = doc.add_paragraph()
            line_content = line_str
            
        # 5. PARSING INLINE (Bold & Italic)
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line_content)
        for token in tokens:
            if not token: continue
            if token.startswith('**') and token.endswith('**') and len(token) > 4:
                run = p.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*') and len(token) > 2:
                run = p.add_run(token[1:-1])
                run.italic = True
            else:
                p.add_run(token)

    # --- TAMBAHAN FOOTER MARKETING ---
    try:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "Generated from rapat.co (formerly tom-stt.com) | The First AI Purpose-Built for Indonesian Transcription and Document Automation"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
    except Exception as e:
        pass 

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def ekstrak_teks_docx_limit(file_bytes, limit=10000):
    import io
    from docx import Document
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        char_count = 0
        for para in doc.paragraphs:
            text = para.text
            if not text.strip(): continue
            full_text.append(text)
            char_count += len(text)
            if char_count >= limit:
                break
        return '\n'.join(full_text)[:limit]
    except Exception as e:
        return ""

def ekstrak_teks_pdf_limit(file_bytes, limit=10000):
    import io
    import PyPDF2
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        full_text = []
        char_count = 0
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
                char_count += len(text)
            if char_count >= limit:
                break
        return '\n'.join(full_text)[:limit]
    except Exception as e:
        return ""

# --- FUNGSI MESIN TRANSKRIP BARU ---
def jalankan_proses_transkrip(audio_to_process, source_name, lang_code):
    st.markdown("---")
    
    status_box = st.empty()
    progress_bar = st.progress(0)
    live_preview_box = st.empty()
    
    full_transcript = []
    
    file_ext = ".wav" if source_name == "rekaman_mic.wav" else (os.path.splitext(source_name)[1] or ".wav")
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
        tmp_file.write(audio_to_process.getvalue())
        input_path = tmp_file.name
        
    # --- FASE 2: THE INTERCEPTOR (AUDIO GATE CERDAS & AMAN) ---
    with st.spinner("🛡️ Menjalankan Front-Gate Validation..."):
        # 1. Ambil durasi riil via FFmpeg
        durasi_detik = get_duration(input_path)
        durasi_menit = math.ceil(durasi_detik / 60)
        
        # 2. Tarik data user terbaru dari Firestore
        u_info = get_user(st.session_state.current_user)
        
        # 3. PENENTUAN HARD-LIMIT PER FILE (Anti-Jebol Server)
        batas_kasta = 20 # Kasta Terendah: Default Freemium 20 Menit
        
        if u_info:
            if u_info.get("role") == "admin":
                batas_kasta = 999999 # Sultan Admin: Bebas hambatan
            elif u_info.get("active_corporate_voucher"):
                batas_kasta = 999999 # 🚀 Kasta B2B Enterprise: Bebas hambatan per file
            elif u_info.get("bank_menit", 0) > 0:
                batas_kasta = 600 # 🚀 Kasta AIO: Maks 10 Jam (600 Menit) per file
            else:
                # Kasta Reguler: Cari tiket dengan batas durasi terbesar di inventori
                max_reguler = 0
                for pkt in u_info.get("inventori", []):
                    if pkt.get("kuota", 0) > 0:
                        max_reguler = max(max_reguler, pkt.get("batas_durasi", 0))
                
                if max_reguler > 0:
                    batas_kasta = max_reguler # Kasta Reguler: Sesuai tiket terbesar
        
        # 4. Filter Kasta (BLOCKIR SEBELUM PROSES MESIN AI)
        if durasi_menit > batas_kasta:
            if os.path.exists(input_path): os.remove(input_path)
            st.error(f"⚠️ **FILE DITOLAK!** Durasi file ({durasi_menit} Menit) melampaui batas PER FILE untuk tier paket Anda (Maks {batas_kasta} Menit).")
            st.info("💡 Demi menjaga kestabilan server AI, silahkan potong audio Anda menjadi beberapa bagian, atau pastikan jenis paket Anda sesuai.")
            st.stop()
            return None

    # --- FASE 2: INJEKSI NYAWA (HARD-CODED LIMIT UNTUK KEAMANAN) ---
    if u_info.get("bank_menit", 0) > 0:
        # Sultan AIO: Ambil dari jatah harian (Default 35)
        st.session_state.sisa_nyawa_dok = u_info.get("fup_dok_harian_limit", 35)
    else:
        # User Reguler (Lite, Starter, dll):
        # Paksa minimal 2 jika di database tidak ada/error, agar tidak langsung habis
        jatah_database = u_info.get("fup_dok_per_file", 2)
        st.session_state.sisa_nyawa_dok = max(2, jatah_database)

    try:
        duration_sec = get_duration(input_path)
        if duration_sec == 0: st.error("Gagal membaca audio."); st.stop()
        
        chunk_len = 59 
        total_chunks = math.ceil(duration_sec / chunk_len)
        
        recognizer = sr.Recognizer()
        recognizer.energy_threshold, recognizer.dynamic_energy_threshold = 300, True 

        status_box.info("⏳ Mempersiapkan mesin transkrip...")

        for i in range(total_chunks):
            start_time = i * chunk_len
            chunk_filename = f"temp_slice_{i}.wav"
            cmd = [ffmpeg_cmd, "-y", "-i", input_path, "-ss", str(start_time), "-t", str(chunk_len), "-filter:a", "volume=3.0", "-ar", "16000", "-ac", "1", chunk_filename]
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            try:
                with sr.AudioFile(chunk_filename) as source:
                    audio_data = recognizer.record(source)
                    text = recognizer.recognize_google(audio_data, language=lang_code)
                    full_transcript.append(text)
            except: full_transcript.append("") 
            finally:
                if os.path.exists(chunk_filename): os.remove(chunk_filename)
            
            progress_percent = int(((i + 1) / total_chunks) * 100)
            progress_bar.progress(progress_percent)
            status_box.caption(f"Sedang memproses... ({progress_percent}%) - MOHON JANGAN TUTUP LAYAR INI!")
            
            partial_text = " ".join(full_transcript)
            st.session_state.transcript = partial_text
            st.session_state.filename = os.path.splitext(source_name)[0]
            
            if st.session_state.logged_in:
                db.collection('users').document(st.session_state.current_user).update({
                    "draft_transcript": partial_text,
                    "draft_filename": st.session_state.filename
                })
            
                live_preview_box.markdown(f"""
                <b style="color: #3498db; font-size: 14px; display: block; margin-bottom: 5px;">Live Preview:</b>
                <div class="no-select" style="background: #F8F9FA; border: 1px solid #DDD; border-radius: 10px; padding: 15px; color: #333; font-size: 13px; line-height: 1.6; max-height: 250px; overflow-y: auto; white-space: pre-wrap; word-wrap: break-word; margin-bottom: 20px;">{partial_text}</div>
                """, unsafe_allow_html=True)

        status_box.success("✔ **Selesai!** Transkrip tersimpan aman. Silahkan klik Tab **🧠 Analisis AI**")
        
        st.session_state.ai_result = "" 
        if st.session_state.logged_in:
            db.collection('users').document(st.session_state.current_user).update({
                "draft_ai_result": "",
                "draft_ai_prefix": ""
            })
        
        st.write("") 
        # 🛡️ TOMBOL DOWNLOAD TXT MENTAH DIHAPUS (GLOBAL SHIELD)

    except Exception as e: 
        status_box.empty()
        st.error(f"Error: {e}")
    finally:
        if os.path.exists(input_path): os.remove(input_path)

# --- BUNGKUS MESIN TRANSKRIP MENJADI FUNGSI AGAR BISA DIPANGGIL DI DALAM TAB ---
def proses_transkrip_audio(audio_to_process, source_name, lang_code):
    st.markdown("---")
    
# 🚀 INJEKSI CSS SEMENTARA (HANYA SELAMA TRANSKRIP BERJALAN)
    stt_css_placeholder = st.empty()
    stt_css_placeholder.markdown("""
    <style>
        /* Animasi berputar eksplisit untuk spinner */
        @keyframes custom-spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        /* Animasi kedap-kedip (pulse) untuk teks */
        @keyframes pulse-text {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.4; }
        }

        /* Timpa Overlay Global menjadi Mini Spinner Kanan Atas */
        [data-testid="stStatusWidget"] {
            top: 20px !important; left: auto !important; right: 20px !important; 
            width: auto !important; height: auto !important; 
            background-color: transparent !important; backdrop-filter: none !important; 
            flex-direction: row !important;
            align-items: center !important; /* Pastikan rata tengah secara vertikal */
        }
        
        [data-testid="stStatusWidget"]::before {
            content: "" !important; /* Wajib ada agar pseudo-element dirender */
            width: 20px !important; height: 20px !important; 
            
            /* Membuat bentuk dan warna spinner */
            border: 3px solid rgba(231, 76, 60, 0.3) !important; /* Warna trek (samar) */
            border-top-color: #e74c3c !important; /* Warna bagian yang berputar (jelas) */
            border-radius: 50% !important; /* Wajib agar bentuknya bulat */
            
            margin-bottom: 0 !important; margin-right: 10px !important; 
            background-color: transparent !important;
            box-shadow: none !important;
            
            /* Terapkan animasi putar */
            animation: custom-spin 1s linear infinite !important;
        }
        
        [data-testid="stStatusWidget"]::after {
            content: "Memproses Audio..." !important; font-size: 13px !important; 
            color: #FFFFFF !important; background-color: rgba(231, 76, 60, 0.9) !important; 
            padding: 5px 12px !important; border-radius: 20px !important; 
            box-shadow: 0 2px 6px rgba(0,0,0,0.15) !important; font-family: 'Plus Jakarta Sans', sans-serif !important;
            
            /* Terapkan animasi kedap-kedip pada teks */
            animation: pulse-text 1.5s ease-in-out infinite !important;
        }
    </style>
    """, unsafe_allow_html=True)

    status_box = st.empty()
    progress_bar = st.progress(0)
    live_preview_box = st.empty()
    
    full_transcript = []
    
    file_ext = ".wav" if source_name == "rekaman_mic.wav" else (os.path.splitext(source_name)[1] or ".wav")
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
        tmp_file.write(audio_to_process.getvalue())
        input_path = tmp_file.name

    # --- FASE 2: THE INTERCEPTOR (AUDIO GATE CERDAS & AMAN) ---
    with st.spinner("🛡️ Menjalankan Front-Gate Validation..."):
        # 1. Ambil durasi riil via FFmpeg
        durasi_detik = get_duration(input_path)
        durasi_menit = math.ceil(durasi_detik / 60)
        
        # 2. Tarik data user terbaru dari Firestore
        u_info = get_user(st.session_state.current_user)
        
        # 3. PENENTUAN HARD-LIMIT PER FILE (Anti-Jebol Server)
        batas_kasta = 20 # Kasta Terendah: Default Freemium 20 Menit
        
        if u_info:
            if u_info.get("role") == "admin":
                batas_kasta = 999999 # Sultan Admin: Bebas hambatan
            elif u_info.get("active_corporate_voucher"):
                batas_kasta = 999999 # 🚀 Kasta B2B Enterprise: Bebas hambatan per file
            elif u_info.get("bank_menit", 0) > 0:
                batas_kasta = 600 # 🚀 Kasta AIO: Maks 10 Jam (600 Menit) per file
            else:
                # Kasta Reguler: Cari tiket dengan batas durasi terbesar di inventori
                max_reguler = 0
                for pkt in u_info.get("inventori", []):
                    if pkt.get("kuota", 0) > 0:
                        max_reguler = max(max_reguler, pkt.get("batas_durasi", 0))
                
                if max_reguler > 0:
                    batas_kasta = max_reguler # Kasta Reguler: Sesuai tiket terbesar
        
        # 4. Filter Kasta (BLOCKIR SEBELUM PROSES MESIN AI)
        if durasi_menit > batas_kasta:
            if os.path.exists(input_path): os.remove(input_path)
            st.error(f"⚠️ **FILE DITOLAK!** Durasi file ({durasi_menit} Menit) melampaui batas PER FILE untuk tier paket Anda (Maks {batas_kasta} Menit).")
            st.info("💡 Demi menjaga kestabilan server AI, silahkan potong audio Anda menjadi beberapa bagian, atau pastikan jenis paket Anda sesuai.")
            st.stop()
            return None

    try:
        duration_sec = get_duration(input_path)
        if duration_sec == 0: st.error("Gagal membaca audio."); st.stop()
        
        # 🛡️ THE DOUBLE SHIELD: CEK DURASI MAKSIMAL & SALDO AIO
        limit_menit = 20
        is_premium = False
        durasi_menit_aktual = math.ceil(duration_sec / 60)
        st.session_state.force_use_reguler_audio = False # 🚀 RESET FLAG FALLBACK
        
        if st.session_state.logged_in:
            usr_cek = get_user(st.session_state.current_user)
            if usr_cek:
                # 🚀 FIX: Akui B2B sebagai Premium agar tidak dicegat batas 20 menit Freemium
                if usr_cek.get("role") == "admin" or len(usr_cek.get("inventori", [])) > 0 or usr_cek.get("active_corporate_voucher"):
                    is_premium = True
                
                # 🚀 GATEWAY CERDAS ALL-IN-ONE & REGULER FALLBACK
                if usr_cek.get("role") != "admin":
                    vid = usr_cek.get("active_corporate_voucher")
                    if vid:
                        v_doc = db.collection('vouchers').document(vid).get()
                        if v_doc.exists:
                            v_data = v_doc.to_dict()
                            sisa_tangki = v_data.get("shared_quota_minutes", 0) - v_data.get("used_quota_minutes", 0)
                            if durasi_menit_aktual > sisa_tangki:
                                status_box.empty()
                                st.error(f"❌ WAKTU INSTANSI TIDAK CUKUP: Audio Anda berdurasi **{durasi_menit_aktual} Menit**, sedangkan sisa tangki instansi Anda hanya **{sisa_tangki} Menit**.")
                                st.warning("Silahkan hubungi PIC Instansi Anda untuk Top-Up kapasitas.")
                                st.stop()
                    else:
                        bank_menit_user = usr_cek.get("bank_menit", 0)
                        
                        # 1. Cek apakah user punya tiket Reguler sebagai cadangan
                        max_durasi_reguler = 0
                        punya_reguler = False
                        for pkt in usr_cek.get("inventori", []):
                            if pkt.get("batas_durasi", 0) != 9999 and pkt.get("kuota", 0) > 0:
                                punya_reguler = True
                                max_durasi_reguler = max(max_durasi_reguler, pkt.get("batas_durasi", 0))

                        if bank_menit_user > 0:
                            if durasi_menit_aktual > bank_menit_user:
                                # AIO KURANG! Coba Fallback ke Reguler
                                if punya_reguler and durasi_menit_aktual <= max_durasi_reguler:
                                    st.session_state.force_use_reguler_audio = True
                                    status_box.empty()
                                    st.info(f"Waktu AIO tidak cukup ({bank_menit_user} mnt). Sistem otomatis mengalihkan pemotongan ke Tiket Reguler Anda.")
                                else:
                                    status_box.empty()
                                    if punya_reguler:
                                        st.error(f"❌ DURASI DITOLAK: Waktu AIO kurang ({bank_menit_user} Mnt), dan durasi audio ({durasi_menit_aktual} Mnt) ini melampaui batas cadangan Paket Reguler Anda (Maks {max_durasi_reguler} Mnt).")
                                    else:
                                        st.error(f"❌ WAKTU AIO TIDAK CUKUP: Audio Anda berdurasi **{durasi_menit_aktual} Menit**, sedangkan sisa Bank Waktu Anda hanya **{bank_menit_user} Menit**.")
                                    st.warning("Silahkan Top-Up Paket Anda terlebih dahulu.")
                                    st.stop()
        
        if not is_premium and durasi_menit_aktual > limit_menit:
            status_box.empty()
            st.error(f"❌ DURASI DITOLAK: Audio Anda berdurasi **{durasi_menit_aktual} Menit**.")
            st.warning(f"Akun Freemium dibatasi maksimal **{limit_menit} Menit**. Silahkan login dan **Beli Paket** untuk memproses audio panjang!")
            st.stop()
        
        # 1. BACA SAKELAR & KASTA USER DARI DATABASE
        sys_config = get_system_config()
        global_use_groq   = sys_config.get("use_groq_stt", False)
        b2b_admin_bypass  = sys_config.get("groq_b2b_admin_bypass", True)
        allowed_packages  = sys_config.get("allowed_packages", [])

        user_info = get_user(st.session_state.current_user) if st.session_state.logged_in else None

        # Menentukan apakah user ini berhak pakai Groq
        # 🚀 Opsi B: B2B/admin bypass bekerja INDEPENDEN dari sakelar retail
        use_groq = False

        if user_info:
            is_admin_role = user_info.get("role") == "admin"
            is_b2b        = bool(user_info.get("active_corporate_voucher"))

            # JALUR 1: B2B/Admin bypass — tidak peduli sakelar retail ON/OFF
            if (is_admin_role or is_b2b) and b2b_admin_bypass:
                use_groq = True

            # JALUR 2: Retail — hanya jika sakelar retail ON dan paket cocok
            elif global_use_groq:
                inventori = user_info.get("inventori", [])
                for pkt in inventori:
                    if pkt['nama'].upper() in [p.upper() for p in allowed_packages]:
                        use_groq = True
                        break
        
        # Cek apakah kunci Groq Whisper tersedia
        active_keys = get_active_keys("Groq Whisper")
        if use_groq and not active_keys:
            use_groq = False
            st.toast("Proses Speech-to-Text dialihkan ke Mesin Transkrip Basic untuk sementara waktu.")

        # ==========================================
        # JALUR 1: MESIN GROQ WHISPER (SAKELAR ON & USER BERHAK)
        # ==========================================
        groq_failed = False

        if use_groq:
            status_box.info("🚀 Mempersiapkan mesin Transkrip HQ...")
            progress_bar.progress(10)
            
            groq_key = active_keys[0]["key"]
            client = Groq(api_key=groq_key)
            
            model_name = sys_config.get("groq_model", "whisper-large-v3")
            short_lang = "id" if lang_code == "id-ID" else "en"
            
            # --- SMART GATEKEEPER V3 (KOMPRESI + PEMOTONGAN CERDAS) ---
            file_size_mb = os.path.getsize(input_path) / (1024 * 1024)
            _, ext = os.path.splitext(input_path)
            ext = ext.lower()
            
            final_audio_path = input_path # Default (Jalur Tol)
            import uuid
            import glob
            
            # TAHAP 1: KOMPRESI JIKA > 15 MB
            if file_size_mb >= 15:
                status_box.caption(f"🗜️ File besar ({file_size_mb:.1f} MB). Melakukan Kompresi...")
                compressed_path = os.path.join(tempfile.gettempdir(), f"compressed_{uuid.uuid4().hex[:6]}.mp3")
                cmd = [ffmpeg_cmd, "-y", "-i", input_path, "-vn", "-ar", "16000", "-ac", "1", "-b:a", "32k", compressed_path]
                subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                final_audio_path = compressed_path
                
            elif ext in ['.opus', '.ogg']:
                status_box.caption(f"Menyesuaikan format audio ke MP3...")
                converted_path = os.path.join(tempfile.gettempdir(), f"converted_{uuid.uuid4().hex[:6]}.mp3")
                cmd = [ffmpeg_cmd, "-y", "-i", input_path, "-vn", "-b:a", "128k", converted_path]
                subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                final_audio_path = converted_path
            else:
                status_box.caption(f"⚡ File {ext} ({file_size_mb:.1f} MB) Aman. Mengirim via Jalur Tol...")
                
            if not os.path.exists(final_audio_path):
                st.error("❌ Gagal memproses audio. Mesin FFmpeg tidak merespons.")
                st.stop()
                
            # TAHAP 2: PEMOTONGAN AUDIO (CHUNKING) JIKA MASIH > 22 MB
            final_size_mb = os.path.getsize(final_audio_path) / (1024 * 1024)
            chunk_files = []
            
            if final_size_mb >= 22:
                status_box.caption(f"✂️ File sangat panjang ({final_size_mb:.1f} MB). Memotong audio menjadi beberapa bagian agar diterima sistem...")
                chunk_prefix = os.path.join(tempfile.gettempdir(), f"chunk_{uuid.uuid4().hex[:6]}_%03d.mp3")
                
                # Memotong file menjadi per 45 menit (2700 detik)
                cmd_split = [ffmpeg_cmd, "-y", "-i", final_audio_path, "-f", "segment", "-segment_time", "2700", "-c", "copy", chunk_prefix]
                subprocess.run(cmd_split, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
                search_pattern = chunk_prefix.replace("%03d", "*")
                chunk_files = sorted(glob.glob(search_pattern))
            else:
                chunk_files = [final_audio_path]
                
            progress_bar.progress(35)
            
            # --- EKSEKUSI API GROQ (MENDUKUNG MULTI-CHUNK) ---
            try:
                hasil_akhir_teks = ""
                total_chunks = len(chunk_files)
                
                for idx, chunk_path in enumerate(chunk_files):
                    status_box.caption(f"☁️ Mengekstrak teks... Bagian {idx+1} dari {total_chunks}")
                    with open(chunk_path, "rb") as audio_file:
                        transcription = client.audio.transcriptions.create(
                            file=(os.path.basename(chunk_path), audio_file.read()),
                            model=model_name,
                            language=short_lang,
                            response_format="text"
                        )
                    # Menjahit teks hasil potongan
                    hasil_akhir_teks += transcription + " "
                    
                    # Animasi Progress Bar
                    prog = 35 + int(((idx + 1) / total_chunks) * 60)
                    progress_bar.progress(prog)
                    
                    # Bersihkan file potongan (chunk) dari server
                    if os.path.exists(chunk_path):
                        os.remove(chunk_path)
                
                hasil_akhir_teks = hasil_akhir_teks.strip()
                progress_bar.progress(100)
                
                # Bersihkan file kompresi utama jika tadi sempat dipecah
                if len(chunk_files) > 1 and final_audio_path != input_path and os.path.exists(final_audio_path):
                    os.remove(final_audio_path)
                elif final_audio_path != input_path and len(chunk_files) == 1 and os.path.exists(final_audio_path):
                    pass # Karena di atas chunk_files sudah di-remove
                    
                # LIVE PREVIEW KHUSUS GROQ
                live_preview_box.markdown(f"""
                <b style="color: #e74c3c; font-size: 14px; display: block; margin-bottom: 5px;">⚡ Hasil Text-to-Speech:</b>
                <div class="no-select" style="background: #F8F9FA; border: 1px solid #DDD; border-radius: 10px; padding: 15px; color: #333; font-size: 13px; line-height: 1.6; max-height: 250px; overflow-y: auto; white-space: pre-wrap; word-wrap: break-word; margin-bottom: 20px;">{hasil_akhir_teks}</div>
                """, unsafe_allow_html=True)
                
                # Hapus / Update API Key Usage
                # Weighted by audio duration: 1 menit audio = +1 increment
                # Sehingga limit Groq Whisper bisa diset dalam satuan "menit audio per hari"
                _groq_menit = max(1, durasi_menit_aktual)
                increment_api_usage(active_keys[0]["id"], active_keys[0]["used"], count=_groq_menit)

            except Exception as _groq_err:
                # ── FALLBACK: Groq gagal → otomatis ke Google chunking ──
                groq_failed = True
                _err_str = str(_groq_err)

                # Cleanup temp files yang mungkin sudah dibuat
                for _cf in chunk_files:
                    if os.path.exists(_cf): os.remove(_cf)
                if final_audio_path != input_path and os.path.exists(final_audio_path):
                    os.remove(final_audio_path)

                # Toast detail error hanya untuk Super Admin Developer
                if st.session_state.get('user_role') == 'admin':
                    st.toast(f"⚠️ Groq Whisper gagal: {_err_str[:120]}. Otomatis dialihkan ke Google.", icon="🔑")

                # Reset UI untuk jalur Google
                progress_bar.progress(0)
                status_box.info("⏳ Mempersiapkan mesin transkrip cadangan...")

        # ==========================================
        # JALUR 2: MESIN GOOGLE (SAKELAR OFF atau GROQ GAGAL)
        # ==========================================
        if not use_groq or groq_failed:
            chunk_len = 59 
            total_chunks = math.ceil(duration_sec / chunk_len)
            recognizer = sr.Recognizer()
            recognizer.energy_threshold, recognizer.dynamic_energy_threshold = 300, True 

            status_box.info("⏳ Mempersiapkan mesin transkrip...")

            for i in range(total_chunks):
                start_time = i * chunk_len
                chunk_filename = f"temp_slice_{i}.wav"
                cmd = [ffmpeg_cmd, "-y", "-i", input_path, "-ss", str(start_time), "-t", str(chunk_len), "-filter:a", "volume=3.0", "-ar", "16000", "-ac", "1", chunk_filename]
                subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
                try:
                    with sr.AudioFile(chunk_filename) as source:
                        audio_data = recognizer.record(source)
                        text = recognizer.recognize_google(audio_data, language=lang_code)
                        full_transcript.append(text)
                except: full_transcript.append("") 
                finally:
                    if os.path.exists(chunk_filename): os.remove(chunk_filename)
                
                progress_percent = int(((i + 1) / total_chunks) * 100)
                progress_bar.progress(progress_percent)
                status_box.caption(f"Sedang memproses... ({progress_percent}%) - MOHON JANGAN TUTUP LAYAR INI!")
                
                partial_text = " ".join(full_transcript)
                st.session_state.transcript = partial_text # Update realtime untuk UI
                
                live_preview_box.markdown(f"""
                <b style="color: #3498db; font-size: 14px; display: block; margin-bottom: 5px;">Live Preview:</b>
                <div class="no-select" style="background: #F8F9FA; border: 1px solid #DDD; border-radius: 10px; padding: 15px; color: #333; font-size: 13px; line-height: 1.6; max-height: 250px; overflow-y: auto; white-space: pre-wrap; word-wrap: break-word; margin-bottom: 20px;">{partial_text}</div>
                """, unsafe_allow_html=True)
                
            hasil_akhir_teks = partial_text

        # --- SAAT PROSES SELESAI (BERLAKU UNTUK KEDUA JALUR) ---
        # 🚀 MENYIAPKAN TEKS STRUK (TANPA DOUBLE DEDUCTION DATABASE)
        teks_struk_aio = ""
        if st.session_state.logged_in:
            usr_akhir = get_user(st.session_state.current_user)
            if usr_akhir and usr_akhir.get("role") != "admin":
                bank_menit_akhir = usr_akhir.get("bank_menit", 0)
                is_fallback = getattr(st.session_state, 'force_use_reguler_audio', False)
                
                # Cek apakah dipotong murni AIO atau beralih (Fallback) ke Reguler
                if bank_menit_akhir > 0 and not is_fallback:
                    new_menit = max(0, bank_menit_akhir - durasi_menit_aktual)
                    teks_struk_aio = f"  \n*Waktu AIO terpotong: **{durasi_menit_aktual} Menit** (Sisa: {new_menit} Menit)*"
                elif is_fallback:
                    teks_struk_aio = f"  \n*Tiket Reguler terpotong (Sisa AIO {bank_menit_akhir} Menit diamankan)*"
                
                # 🚀 STRICT ORIGIN FUP: FUP DIBERIKAN SESUAI TIKET YANG DIPOTONG
                # Jika dipotong murni pakai Bank Menit AIO (Bukan Fallback)
                if usr_akhir.get("bank_menit", 0) > 0 and not is_fallback:
                    st.session_state.sisa_nyawa_dok = usr_akhir.get("fup_dok_harian_limit", 35)
                else:
                    # Jika beralih ke Reguler (Fallback) atau murni Reguler
                    max_fup = 2
                    for pkt in usr_akhir.get("inventori", []):
                        p_name = pkt.get("nama", "").upper()
                        # Pastikan kita hanya membaca kasta tiket Reguler yang masih ada kuotanya
                        if "AIO" not in p_name and pkt.get("kuota", 0) > 0:
                            if "ENTERPRISE" in p_name: max_fup = max(max_fup, 20)
                            elif "VIP" in p_name: max_fup = max(max_fup, 12)
                            elif "EKSEKUTIF" in p_name: max_fup = max(max_fup, 8)
                            elif "STARTER" in p_name: max_fup = max(max_fup, 4)
                            elif "LITE" in p_name: max_fup = max(max_fup, 2)
                    st.session_state.sisa_nyawa_dok = max_fup

        status_box.success(f"**Selesai!** Transkrip tersimpan aman.\n\nDurasi Asli Audio: **{durasi_menit_aktual} Menit**{teks_struk_aio}\n\nLanjutkan ke **🧠 Analisis AI**")
        
        # Simpan durasi kotor ke memori agar bisa dibaca di Tab 4
        st.session_state.durasi_audio_kotor = durasi_menit_aktual
        
        st.session_state.transcript = hasil_akhir_teks
        st.session_state.filename = os.path.splitext(source_name)[0]
        st.session_state.ai_result = "" 
        
        # --- PERBAIKAN: EKSEKUSI PEMOTONGAN KUOTA AUDIO ---
        if st.session_state.logged_in:
            u_doc = db.collection('users').document(st.session_state.current_user)
            u_info = u_doc.get().to_dict()
            
            # Ambil kembali durasi menit kotor yang sudah disiapkan
            durasi_menit = st.session_state.get('durasi_audio_kotor', 1)
            
            # 🚀 LOGIKA PEMOTONGAN CERDAS (SUPPORT FALLBACK & B2B)
            is_fallback_reguler = getattr(st.session_state, 'force_use_reguler_audio', False)
            vid = u_info.get("active_corporate_voucher")
            
            if vid:
                # 1. User B2B Enterprise: Potong dari Tangki Instansi & Catat Analitik Staf
                v_ref = db.collection('vouchers').document(vid)
                v_doc_snap = v_ref.get()
                if v_doc_snap.exists:
                    v_data = v_doc_snap.to_dict()
                    curr_staff = v_data.get("staff_usage", {})
                    user_email_key = st.session_state.current_user
                    
                    # 🚀 FIX: Update dict di Python lalu set kembali, hindari bug dot notation
                    if user_email_key not in curr_staff:
                        curr_staff[user_email_key] = {"minutes_used": 0, "docs_generated": 0}
                    curr_staff[user_email_key]["minutes_used"] += durasi_menit
                    curr_staff[user_email_key]["docs_generated"] += 1
                    
                    v_ref.update({
                        "used_quota_minutes": firestore.Increment(durasi_menit),
                        "total_documents_generated": firestore.Increment(1),
                        "staff_usage": curr_staff
                    })
                st.toast(f"Tangki Instansi terpotong {durasi_menit} Menit", icon="⏳")
            elif u_info.get("bank_menit", 0) > 0 and not is_fallback_reguler:
                # 2. User AIO Normal: Potong saldo bank menit
                new_bank = max(0, u_info["bank_menit"] - durasi_menit)
                u_doc.update({"bank_menit": new_bank})
                st.toast(f"Saldo Paket AIO terpotong {durasi_menit} Menit", icon="⏳")
            else:
                # 3. User Reguler ATAU Fallback AIO: Potong 1 Tiket Reguler
                inv = u_info.get("inventori", [])
                idx_to_cut = -1
                # Cari index paket reguler pertama (Abaikan paket AIO yang batasnya 9999)
                for i, pkt in enumerate(inv):
                    if pkt.get('batas_durasi', 0) != 9999 and pkt.get('kuota', 0) > 0:
                        idx_to_cut = i
                        break
                
                if idx_to_cut != -1:
                    inv[idx_to_cut]['kuota'] -= 1
                    if inv[idx_to_cut]['kuota'] <= 0:
                        inv.pop(idx_to_cut)
                    u_doc.update({"inventori": inv})
                    
                    if is_fallback_reguler:
                        st.toast("🎟️ 1 Tiket Reguler terpotong (Efek Fallback AIO)!", icon="✔")
                    else:
                        st.toast("🎟️ 1 Tiket Transkrip Audio terpotong!", icon="✔")

            # 3. Simpan Draft Transkrip terakhir ke Firestore
            u_doc.update({
                "draft_transcript": hasil_akhir_teks,
                "draft_filename": st.session_state.filename,
                "draft_ai_result": "",
                "draft_ai_prefix": "",
                "is_text_upload": False
            })
            
            # ✅ Invalidasi cache setelah semua write selesai.
            # Memastikan logika FUP di bawah & get_user() manapun
            # yang dipanggil setelahnya membaca data Firestore yang sudah dipotong.
            invalidate_user_cache(st.session_state.current_user)
            
            # --- PERBAIKAN LOGIKA FUP: STRICT ORIGIN (SESUAI TIKET YG DIPOTONG) ---
            is_fallback = getattr(st.session_state, 'force_use_reguler_audio', False)
            vid = u_info.get("active_corporate_voucher")
            
            if vid:
                # 1. Jalur B2B Enterprise — FUP sesuai harian limit voucher (anti-spam, bukan unlimited)
                st.session_state.sisa_nyawa_dok = u_info.get("fup_dok_harian_limit", 35)
                st.session_state.is_using_aio = False
            elif u_info.get("bank_menit", 0) > 0 and not is_fallback:
                # 2. Jalur Sultan (Murni AIO)
                st.session_state.sisa_nyawa_dok = u_info.get("fup_dok_harian_limit", 35)
                st.session_state.is_using_aio = True
            else:
                # 3. Jalur Reguler (Tiket Reguler Terpotong)
                max_fup = 2
                for pkt in u_info.get("inventori", []):
                    p_name = pkt.get("nama", "").upper()
                    if "AIO" not in p_name and pkt.get("kuota", 0) > 0:
                        if "ENTERPRISE" in p_name: max_fup = max(max_fup, 20)
                        elif "VIP" in p_name: max_fup = max(max_fup, 12)
                        elif "EKSEKUTIF" in p_name: max_fup = max(max_fup, 8)
                        elif "STARTER" in p_name: max_fup = max(max_fup, 4)
                        elif "LITE" in p_name: max_fup = max(max_fup, 2)
                st.session_state.sisa_nyawa_dok = max_fup
                st.session_state.is_using_aio = False
                    
        st.write("")
        
        # 🔥 FITUR BARU: TOMBOL PINDAH TAB OTOMATIS (JAVASCRIPT INJECTION)
        # Tombol ini dibuat menggunakan HTML/JS agar saat diklik, ia akan mencari 
        # Tab 'Analisis AI' di sistem Streamlit dan berpindah seketika (Instan).
        # 🔧 MIGRASI: components.html → st.html(unsafe_allow_javascript=True)
        # window.parent.document → document, window.parent.scrollTo → window.scrollTo
        # 🐛 FIX: Inline onclick="" di-strip oleh Streamlit st.html() sanitization.
        # Solusi: pakai id + assign onclick via <script> IIFE (pattern sama dengan token card LOKASI 5).
        btn_html = """
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@700&display=swap');
            .btn-switch-stt {
                background-color: #000000; color: #FFFFFF; font-family: 'Plus Jakarta Sans', sans-serif;
                border: none; padding: 14px 20px; font-size: 16px; font-weight: 700;
                border-radius: 10px; width: 100%; cursor: pointer; transition: all 0.2s;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1); display: block; box-sizing: border-box;
            }
            .btn-switch-stt:hover { background-color: #333333; transform: translateY(-2px); }
        </style>
        <button class="btn-switch-stt" id="rapatco-btn-switch-stt">🧠 Lanjut ke Analisis AI</button>
        <script>
        (function() {
            var btn = document.getElementById("rapatco-btn-switch-stt");
            if (!btn) return;
            btn.onclick = function() {
                var tabs = document.querySelectorAll('button[data-baseweb="tab"]');
                var targetTab = Array.from(tabs).find(function(tab) {
                    return tab.innerText.includes('Analisis AI');
                });
                if (targetTab) {
                    targetTab.click();
                    window.scrollTo({top: 0, behavior: 'smooth'});
                }
            };
        })();
        </script>
        """
        st.html(btn_html, unsafe_allow_javascript=True)

    except Exception as e:
        status_box.empty()
        st.error(f"Error: {str(e)}")
    finally:
        if os.path.exists(input_path): os.remove(input_path)
        # 🚀 CABUT CSS SEMENTARA AGAR KEMBALI KE OVERLAY GLOBAL
        stt_css_placeholder.empty()