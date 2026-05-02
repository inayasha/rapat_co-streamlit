"""
engine_template.py — AI Custom Template Engine

Tiga komponen utama:
  1. ekstrak_struktur_docx()     — Smart extractor: baca teks, style, tabel, header/footer
  2. build_prompt_template()     — Bangun prompt dengan structural brief untuk AI
  3. create_docx_from_markers()  — Renderer: ubah marker output AI jadi elemen .docx nyata
  4. render_custom_template_ui() — UI + billing logic, dipanggil dari app.py
"""

import io
import re
import time
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from firebase_admin import firestore
import google.generativeai as genai

from database import (
    db, get_user, get_active_keys, get_system_config,
    increment_api_usage, invalidate_user_cache
)


# ============================================================
# 1. SMART EXTRACTOR
# ============================================================

def ekstrak_struktur_docx(file_bytes: bytes, limit: int = 10000) -> str:
    """
    Membaca .docx dan menghasilkan 'structural brief' yang kaya untuk AI.

    PENTING: Format output menggunakan notasi '>> ' (double-arrow) yang
    bersifat DESKRIPTIF — bukan marker yang boleh direproduksi AI ke output.

    Urutan prioritas:
      A. Header & footer Word (jika ada)
      B. Semua tabel (kop manual, tanda tangan, data)
      C. Paragraf dengan metadata style (heading level, alignment, bold)
      D. Section margins sebagai petunjuk layout
    """
    try:
        doc        = Document(io.BytesIO(file_bytes))
        parts      = []
        char_count = 0

        def _add(text: str) -> bool:
            nonlocal char_count
            if char_count >= limit:
                return False
            parts.append(text)
            char_count += len(text)
            return True

        # ── A. HEADER & FOOTER WORD ──────────────────────────────────
        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                t = hdr_para.text.strip()
                if t:
                    if not _add(f">> Header Word: {t}\n"):
                        break
            for ftr_para in section.footer.paragraphs:
                t = ftr_para.text.strip()
                if t:
                    if not _add(f">> Footer Word: {t}\n"):
                        break

        # ── B. SEMUA TABEL ───────────────────────────────────────────
        for tbl_idx, table in enumerate(doc.tables):
            if char_count >= limit:
                break

            rows_data = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                # Deduplicate merged cells yang muncul berulang
                unique_cells = []
                prev = None
                for c in cells:
                    if c != prev:
                        unique_cells.append(c)
                    prev = c
                row_text = " | ".join(filter(None, unique_cells))
                if row_text.strip():
                    rows_data.append(row_text)

            if rows_data:
                tbl_label = f">> Tabel {tbl_idx + 1} ({len(rows_data)} baris):\n"
                if not _add(tbl_label):
                    break
                for row_text in rows_data:
                    if not _add(f"   {row_text}\n"):
                        break

        # ── C. PARAGRAF DENGAN METADATA STYLE ────────────────────────
        ALIGN_MAP = {
            WD_ALIGN_PARAGRAPH.CENTER:  "rata tengah",
            WD_ALIGN_PARAGRAPH.RIGHT:   "rata kanan",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "rata kanan-kiri",
            WD_ALIGN_PARAGRAPH.LEFT:    "rata kiri",
        }

        for para in doc.paragraphs:
            if char_count >= limit:
                break

            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            align      = ALIGN_MAP.get(para.alignment, "rata kiri")

            # Deteksi heading level
            heading_match = re.match(r'Heading (\d)', style_name)
            if heading_match:
                lvl = heading_match.group(1)
                if not _add(f">> Judul Level {lvl} ({align}): {text}\n"):
                    break
                continue

            # Deteksi bold (cek run pertama yang punya teks)
            is_bold = any(run.bold for run in para.runs if run.text.strip())

            # Deteksi penomoran dari XML
            is_numbered = para._element.find(qn('w:numPr')) is not None

            # Susun deskripsi style dalam bahasa natural
            style_desc_parts = []
            if is_bold:       style_desc_parts.append("tebal")
            if align != "rata kiri": style_desc_parts.append(align)
            if is_numbered:   style_desc_parts.append("item list")

            if style_desc_parts:
                style_desc = ", ".join(style_desc_parts)
                if not _add(f">> Paragraf ({style_desc}): {text}\n"):
                    break
            else:
                if not _add(f">> Paragraf: {text}\n"):
                    break

        # ── D. MARGIN SEBAGAI PETUNJUK LAYOUT ────────────────────────
        try:
            sec = doc.sections[0]
            margin_info = (
                f">> Layout halaman: margin atas={sec.top_margin.cm:.1f}cm, "
                f"bawah={sec.bottom_margin.cm:.1f}cm, "
                f"kiri={sec.left_margin.cm:.1f}cm, "
                f"kanan={sec.right_margin.cm:.1f}cm\n"
            )
            _add(margin_info)
        except Exception:
            pass

        return "".join(parts)[:limit]

    except Exception:
        return ""


# ============================================================
# 2. PROMPT BUILDER
# ============================================================

def build_prompt_template(structural_brief: str, nama_dok: str) -> str:
    """
    Membangun prompt Custom Template yang menginstruksikan AI untuk:
    1. Membaca structural brief (notasi '>>' = DESKRIPSI, bukan template)
    2. Menghasilkan output dengan OUTPUT MARKER terstruktur
    """
    return f"""Anda adalah Asisten AI Pembuat Dokumen Profesional Indonesia.
Tugas: susun [{nama_dok.upper()}] baru dari [TRANSKRIP SUMBER], dengan meniru \
PERSIS struktur, hierarki, dan gaya bahasa dari [DOKUMEN ACUAN].

=== PERINGATAN KERAS — BACA SEBELUM MEMPROSES ===

Bagian [DOKUMEN ACUAN] di bawah menggunakan notasi deskripsi internal saya \
yang diawali dengan ">>" (contoh: ">> Paragraf (tebal, rata tengah): JUDUL").
Notasi ">>" adalah CATATAN SAYA UNTUK MENDESKRIPSIKAN template — bukan bagian \
dari dokumen yang harus Anda tulis ulang.

LARANGAN KERAS:
- JANGAN tulis ">>" dalam output Anda
- JANGAN tulis ">> Paragraf", ">> Tabel", ">> Judul Level" atau notasi deskripsi apapun
- JANGAN copy-paste teks dari bagian [DOKUMEN ACUAN] ke output
- Output Anda HANYA berisi konten dokumen baru + OUTPUT MARKER yang didefinisikan di bawah

=== OUTPUT MARKER (WAJIB DIGUNAKAN UNTUK ELEMEN VISUAL) ===

Gunakan marker berikut HANYA untuk elemen yang TIDAK bisa ditulis sebagai teks biasa:

[PLACEHOLDER:LOGO]
  Tulis ini jika dokumen acuan mengindikasikan ada logo/lambang instansi.

[PLACEHOLDER:KOP | teks: Nama Instansi dan alamat]
  Tulis ini untuk area kop surat teks. Isi 'teks' dengan informasi kop dari acuan,
  atau tulis "[Nama Instansi]" jika tidak ada di transkrip.

[GARIS:TEBAL]
  Garis horizontal tebal (biasanya di bawah kop surat).

[GARIS:TIPIS]
  Garis horizontal tipis (pemisah antar seksi).

[JUDUL | align:center | style:bold | teks: JUDUL DOKUMEN]
  Judul utama dokumen. Sesuaikan align (center/left/right) dengan dokumen acuan.

[TABEL_TTD | kolom: Label Kiri, Label Kanan]
  Kolom tanda tangan. Buat sesuai jumlah kolom di dokumen acuan.

[FIELD | label: Nomor | nilai: [isi dari transkrip atau placeholder]]
  Field formulir: Nomor, Tanggal, Perihal, Kepada, dll.

Untuk konten teks biasa (paragraf, heading, bullet, tabel data):
  Tulis langsung sebagai teks — JANGAN gunakan notasi ">>" atau tag apapun selain
  OUTPUT MARKER di atas.

=== INSTRUKSI KONTEN ===

1. KLONING STRUKTUR: Tiru PERSIS hierarki dan penomoran (I/A/1/a atau lainnya)
   dari Dokumen Acuan. Jika acuan pakai "I. PENDAHULUAN" maka output harus sama.

2. ISOLASI KONTEN: JANGAN menyalin nama, tanggal, lokasi, nominal, atau fakta
   spesifik dari Dokumen Acuan. Semua fakta MURNI dari Transkrip Sumber.

3. PLACEHOLDER DATA: Field yang tidak ada di Transkrip → tulis [Sesuaikan manual].
   JANGAN hapus field tersebut dan JANGAN mengarang fakta.

4. DETEKSI ELEMEN VISUAL DARI DESKRIPSI ACUAN:
   - ">> Tabel 1" di baris pertama dengan konten nama instansi/alamat
     → ini kop surat manual → gunakan [PLACEHOLDER:LOGO] + [PLACEHOLDER:KOP]
   - ">> Tabel N" di akhir dokumen berisi kata "Mengetahui"/"Menyetujui"/"Notulis"
     → ini kolom tanda tangan → gunakan [TABEL_TTD]
   - Baris yang isinya "---" atau "___" berulang → gunakan [GARIS:TEBAL/TIPIS]

5. ANTI BASA-BASI: Karakter PERTAMA output HARUS langsung berupa marker atau
   isi dokumen. DILARANG menyapa, memberi pengantar, atau penutup.
   JANGAN gunakan blok kode markdown (jangan gunakan ```).

=== DOKUMEN ACUAN (BACA — JANGAN COPY NOTASINYA) ===
{structural_brief}
"""


# ============================================================
# 3. DOCX RENDERER (marker → elemen Word nyata)
# ============================================================

def _add_horizontal_line(paragraph, thickness_pt: float = 1.0, color: str = "000000"):
    """Tambahkan garis horizontal di bawah paragraf via XML border."""
    pPr    = paragraph._element.get_or_add_pPr()
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(int(thickness_pt * 8)))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _placeholder_box(doc: Document, label: str, hint: str):
    """Buat kotak placeholder berformat (1x1 tabel, border abu-abu)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    tc_pr    = cell._tc.get_or_add_tcPr()
    tc_bords = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'AAAAAA')
        tc_bords.append(b)
    tc_pr.append(tc_bords)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F5F5F5')
    tc_pr.append(shd)

    para       = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label  = para.add_run(f"[ {label} ]")
    run_label.bold      = True
    run_label.italic    = True
    run_label.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run_label.font.size = Pt(10)

    if hint:
        run_hint = para.add_run(f"\n{hint}")
        run_hint.italic         = True
        run_hint.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run_hint.font.size      = Pt(8.5)

    tc_mar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    '120')
        m.set(qn('w:type'), 'dxa')
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    doc.add_paragraph()


def _add_signature_table(doc: Document, labels: list):
    """Buat tabel tanda tangan N kolom tanpa border, siap diisi."""
    n = len(labels)
    if n == 0:
        return

    table           = doc.add_table(rows=4, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, label in enumerate(labels):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label)
        run.bold      = True
        run.font.size = Pt(10.5)

    for row_idx in [1, 2]:
        for i in range(n):
            cell = table.cell(row_idx, i)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(" ").font.size = Pt(10.5)

    for i in range(n):
        cell = table.cell(3, i)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("(______________________)").font.size = Pt(10.5)

    doc.add_paragraph()


def create_docx_from_markers(ai_text: str, title: str) -> bytes:
    """
    Mengubah output AI (teks + OUTPUT MARKER) menjadi file .docx.

    OUTPUT MARKER yang dikenali:
      [PLACEHOLDER:LOGO]
      [PLACEHOLDER:KOP | teks: ...]
      [GARIS:TEBAL]
      [GARIS:TIPIS]
      [JUDUL | align:... | style:... | teks: ...]
      [TABEL_TTD | kolom: Label1, Label2, ...]
      [FIELD | label: ... | nilai: ...]
      Heading markdown: ## Judul
      Bold/italic markdown: **teks** / *teks*
      Tabel markdown: | col1 | col2 |
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    style          = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Set spacing semua style: 0pt before/after, single line spacing
    # Heading mendapat space_before 6pt agar ada pemisah visual
    for s in doc.styles:
        try:
            pf = s.paragraph_format
            is_heading = 'Heading' in s.name
            pf.space_before      = Pt(6) if is_heading else Pt(0)
            pf.space_after       = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except Exception:
            pass

    in_table_md = False
    table_md    = None

    ALIGN_MAP_STR = {
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'right':   WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
    }

    def _parse_inline(para, text: str):
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**') and len(token) > 4:
                run      = para.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*') and len(token) > 2:
                run        = para.add_run(token[1:-1])
                run.italic = True
            else:
                para.add_run(token)

    # Bersihkan sisa notasi '>>' jika AI masih menyelipkannya (safety net)
    def _strip_arrow_notation(text: str) -> str:
        return re.sub(r'^>>\s*(?:Paragraf|Tabel|Judul Level|Header|Footer|Layout)\s*(?:\([^)]*\))?\s*:?\s*', '', text).strip()

    lines = ai_text.split('\n')

    for line in lines:
        # Safety net: hapus notasi '>>' jika masih ada
        if line.strip().startswith('>>'):
            cleaned = _strip_arrow_notation(line.strip())
            if not cleaned:
                continue
            # Render sebagai paragraf biasa tanpa tag
            p = doc.add_paragraph()
            _parse_inline(p, cleaned)
            continue

        raw = line.strip()
        if not raw:
            in_table_md = False
            doc.add_paragraph()
            continue

        # ── [PLACEHOLDER:LOGO] ────────────────────────────────────
        if re.match(r'\[PLACEHOLDER:LOGO\]', raw, re.IGNORECASE):
            _placeholder_box(doc,
                label="LOGO / LAMBANG INSTANSI",
                hint="Letakkan logo instansi Anda di sini (di Word: Insert → Picture)"
            )
            continue

        # ── [PLACEHOLDER:KOP | teks: ...] ────────────────────────
        m = re.match(r'\[PLACEHOLDER:KOP\s*\|?\s*teks:\s*(.*)\]', raw, re.IGNORECASE)
        if m:
            kop_text = m.group(1).strip()
            _placeholder_box(doc,
                label="KOP SURAT",
                hint=kop_text if kop_text else "Nama instansi, alamat, nomor telepon"
            )
            continue

        # ── [GARIS:TEBAL] ─────────────────────────────────────────
        if re.match(r'\[GARIS:TEBAL\]', raw, re.IGNORECASE):
            p = doc.add_paragraph()
            _add_horizontal_line(p, thickness_pt=2.0, color="000000")
            continue

        # ── [GARIS:TIPIS] ─────────────────────────────────────────
        if re.match(r'\[GARIS:TIPIS\]', raw, re.IGNORECASE):
            p = doc.add_paragraph()
            _add_horizontal_line(p, thickness_pt=0.5, color="555555")
            continue

        # ── [JUDUL | align:... | teks: ...] ──────────────────────
        m = re.match(r'\[JUDUL\s*\|?\s*(.*?)\s*\|?\s*teks:\s*(.*)\]', raw, re.IGNORECASE)
        if m:
            meta_str  = m.group(1).strip()
            judul_txt = m.group(2).strip()
            align_str = "center"
            is_bold   = True
            for part in meta_str.split('|'):
                part = part.strip().lower()
                if part.startswith('align:'):
                    align_str = part.replace('align:', '').strip()

            p           = doc.add_paragraph()
            p.alignment = ALIGN_MAP_STR.get(align_str, WD_ALIGN_PARAGRAPH.CENTER)
            run         = p.add_run(judul_txt)
            run.bold      = is_bold
            run.font.size = Pt(14)
            continue

        # ── [TABEL_TTD | kolom: ...] ──────────────────────────────
        m = re.match(r'\[TABEL_TTD\s*\|?\s*kolom:\s*(.*)\]', raw, re.IGNORECASE)
        if m:
            labels = [lb.strip() for lb in m.group(1).split(',') if lb.strip()]
            if not labels:
                labels = ["Mengetahui", "Notulis"]
            _add_signature_table(doc, labels)
            continue

        # ── [FIELD | label: ... | nilai: ...] ────────────────────
        m = re.match(r'\[FIELD\s*\|?\s*label:\s*(.*?)\s*\|?\s*nilai:\s*(.*)\]', raw, re.IGNORECASE)
        if m:
            label_f = m.group(1).strip()
            nilai_f = m.group(2).strip()
            p       = doc.add_paragraph()
            run_l   = p.add_run(f"{label_f}\t: ")
            run_l.bold = True
            p.add_run(nilai_f)
            continue

        # ── TABEL MARKDOWN | col | col | ──────────────────────────
        # Handle dua format:
        #   Format A (standar)  : | col1 | col2 | col3 |
        #   Format B (AI style) : col1 | col2 | col3    (tanpa leading/trailing pipe)
        if raw.count('|') >= 2:
            # Normalize: buang leading/trailing pipe jika ada, lalu split
            normalized = raw
            if normalized.startswith('|'):
                normalized = normalized[1:]
            if normalized.endswith('|'):
                normalized = normalized[:-1]
            cells = [c.strip() for c in normalized.split('|')]
            # Lewati baris separator (--- atau :---:)
            if all(re.match(r'^[-:\s]+$', c) for c in cells if c):
                continue
            if not in_table_md:
                in_table_md = True
                table_md    = doc.add_table(rows=1, cols=len(cells))
                table_md.style = 'Table Grid'
                hdr = table_md.rows[0].cells
                for i, val in enumerate(cells):
                    if i < len(hdr):
                        clean = val.replace('**', '').replace('*', '')
                        hdr[i].text = clean
                        if hdr[i].paragraphs and hdr[i].paragraphs[0].runs:
                            hdr[i].paragraphs[0].runs[0].bold = True
            else:
                row_cells = table_md.add_row().cells
                for i, val in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = val.replace('**', '').replace('*', '')
            continue
        else:
            in_table_md = False

        # ── HEADING MARKDOWN ## ────────────────────────────────────
        hm = re.match(r'^(#{1,4})\s+(.*)', raw)
        if hm:
            lvl = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2), level=lvl)
            continue

        # ── GARIS MARKDOWN --- ─────────────────────────────────────
        if re.match(r'^-{3,}$', raw):
            p = doc.add_paragraph()
            _add_horizontal_line(p, thickness_pt=0.5, color="888888")
            continue

        # ── BULLET ────────────────────────────────────────────────
        bm = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bm:
            indent = len(bm.group(1))
            p = doc.add_paragraph(
                style='List Bullet' if indent < 2 else 'List Bullet 2'
            )
            _parse_inline(p, bm.group(2))
            continue

        # ── NUMBERED ──────────────────────────────────────────────
        nm = re.match(r'^(\s*)([A-Za-z0-9]+[.)]\s+)(.*)', line)
        if nm:
            p      = doc.add_paragraph()
            indent = len(nm.group(1))
            if indent > 0:
                p.paragraph_format.left_indent = Pt(18)
            _parse_inline(p, nm.group(2) + nm.group(3))
            continue

        # ── PARAGRAF BIASA ─────────────────────────────────────────
        p = doc.add_paragraph()
        _parse_inline(p, raw)

    # Footer marketing
    try:
        footer = doc.sections[0].footer
        fp     = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = "Generated from rapat.co (formerly tom-stt.com) | TEMAN RAPAT / TOM'STT AI"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in fp.runs:
            run.font.size      = Pt(8)
            run.font.italic    = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    except Exception:
        pass

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ============================================================
# 4. UI + BILLING LOGIC — dipanggil dari app.py
# ============================================================

def render_custom_template_ui(user_info: dict, sys_config: dict):
    """
    Render seluruh UI 🎨 AI Custom Template termasuk billing.
    Dipanggil dari dalam blok with wadah_tombol: di app.py.
    """
    st.info(
        "💡 **Bring Your Own Template (BYOT):** AI akan menulis dokumen "
        "menggunakan fakta dari transkrip, namun meniru tata letak, gaya bahasa, "
        "dan elemen visual (kop surat, garis, kolom tanda tangan) dari dokumen acuan Anda."
    )
    st.info(
        "📱 **Tips Pengguna HP:** Pastikan Anda sudah mengetahui letak dokumen "
        "acuan (.docx) sebelum menekan *Upload* agar koneksi tidak terputus."
    )

    # ── Billing state ─────────────────────────────────────────────
    kuota_custom  = 2
    saldo_user    = 0
    is_b2b_custom = False
    vid_b2b       = None
    is_admin      = st.session_state.user_role == "admin"

    if not is_admin and user_info:
        kuota_custom  = user_info.get("kuota_custom_gratis", 2)
        saldo_user    = user_info.get("saldo", 0)
        vid_b2b       = user_info.get("active_corporate_voucher")
        is_b2b_custom = bool(vid_b2b)

    # ── Upload — hanya .docx ──────────────────────────────────────
    st.markdown(
        "<p style='font-size:14px; margin-bottom:-10px;'>1. Upload Dokumen Acuan (.docx)</p>",
        unsafe_allow_html=True
    )
    uploaded_template = st.file_uploader(
        "",
        type=["docx"],
        accept_multiple_files=False,
        help=(
            "Sistem membaca struktur, tabel, heading, dan margin dokumen secara cerdas "
            "(maks. 10.000 karakter). Elemen visual seperti logo akan dibuatkan "
            "kotak placeholder di output untuk diisi manual di Word."
        ),
        label_visibility="collapsed",
        key="custom_template_uploader"
    )

    jenis_dok_custom = st.text_input(
        "2. Jenis Dokumen yang Ingin Dibuat (Opsional)",
        placeholder="Contoh: Laporan SPPD, Notulen Rapat, Putusan Sidang, dll."
    ).strip()

    if not uploaded_template:
        return

    nama_file_temp = uploaded_template.name
    is_revisi      = (nama_file_temp == st.session_state.get('custom_template_last_file', ''))

    # ── Label tombol & tarif ──────────────────────────────────────
    if is_admin:
        teks_tombol  = "✨ Generate Custom Format (Admin: Gratis)"
        harga_custom = 0
        menit_b2b    = 0
    elif is_b2b_custom:
        menit_b2b    = 5 if is_revisi else 15
        label_revisi = "Revisi" if is_revisi else "Buat Baru"
        teks_tombol  = f"✨ Generate Custom Format (B2B {label_revisi}: {menit_b2b} Menit)"
        harga_custom = 0
    elif kuota_custom > 0:
        teks_tombol  = f"✨ Generate Custom Format (Gratis: Sisa {kuota_custom}x)"
        harga_custom = 0
        menit_b2b    = 0
    elif is_revisi:
        teks_tombol  = "🔄 Generate Ulang / Revisi (Tarif: Rp 2.500)"
        harga_custom = 2_500
        menit_b2b    = 0
    else:
        teks_tombol  = "✨ Generate Custom Format (Tarif: Rp 10.000)"
        harga_custom = 10_000
        menit_b2b    = 0

    if not st.button(teks_tombol, width='stretch', type="primary",
                     key="btn_custom_template_generate"):
        return

    # ── Validasi keuangan ─────────────────────────────────────────
    if is_b2b_custom:
        v_snap = db.collection('vouchers').document(vid_b2b).get()
        if v_snap.exists:
            v_data      = v_snap.to_dict()
            sisa_tangki = (v_data.get("shared_quota_minutes", 0)
                           - v_data.get("used_quota_minutes", 0))
            if sisa_tangki < menit_b2b:
                st.error(f"❌ **KUOTA INSTANSI HABIS!** Butuh {menit_b2b} Menit untuk format ini.")
                st.stop()
        else:
            st.error("❌ Data lisensi instansi tidak ditemukan.")
            st.stop()
    elif harga_custom > 0 and saldo_user < harga_custom:
        st.error(f"❌ **SALDO TIDAK CUKUP!** Tarif layanan ini Rp {harga_custom:,}.")
        st.warning("💡 Silahkan Top-Up Saldo Utama Anda di menu samping.")
        st.stop()

    # ── Ekstrak struktur dokumen (smart extractor) ────────────────
    with st.spinner("🔍 Membaca struktur dokumen acuan..."):
        file_bytes       = uploaded_template.read()
        structural_brief = ekstrak_struktur_docx(file_bytes, limit=10_000)

        if not structural_brief.strip():
            st.error(
                "❌ Gagal membaca dokumen. Pastikan file .docx tidak kosong, "
                "tidak diproteksi password, dan bukan format .doc lama."
            )
            st.stop()

    # ── Eksekusi pembayaran ───────────────────────────────────────
    if not is_admin:
        u_doc = db.collection('users').document(st.session_state.current_user)
        if is_b2b_custom:
            db.collection('vouchers').document(vid_b2b).update({
                "used_quota_minutes": firestore.Increment(menit_b2b)
            })
            st.toast(f"🏛️ Tangki Instansi terpotong {menit_b2b} Menit.", icon="✔")
        else:
            if kuota_custom > 0:
                u_doc.update({"kuota_custom_gratis": kuota_custom - 1})
                st.toast(f"🎁 Kuota Gratis terpakai. Sisa: {kuota_custom - 1}x", icon="✔")
            else:
                u_doc.update({"saldo": saldo_user - harga_custom})
                st.toast(f"Saldo terpotong Rp {harga_custom:,}", icon="💳")

        if 'temp_user_data' in st.session_state:
            del st.session_state['temp_user_data']

    st.session_state.custom_template_last_file = nama_file_temp

    # ── Bangun prompt & panggil Gemini ────────────────────────────
    nama_dok      = jenis_dok_custom if jenis_dok_custom else "dokumen"
    prompt_custom = build_prompt_template(structural_brief, nama_dok)

    loading = st.empty()
    loading.markdown("""
    <style>
    .loading-screen {
        position:fixed; top:0; left:0; width:100vw; height:100vh;
        background-color:rgba(255,255,255,0.92);
        display:flex; flex-direction:column; justify-content:center; align-items:center;
        z-index:999999; backdrop-filter:blur(8px);
    }
    .spinner-large {
        width:50px; height:50px; border:5px solid #F0F2F6;
        border-top:5px solid #e74c3c; border-radius:50%;
        animation:spin-large 1s linear infinite; margin-bottom:15px;
    }
    @keyframes spin-large{0%{transform:rotate(0deg)}100%{transform:rotate(360deg)}}
    .loading-title{font-size:17px;font-weight:600;color:#333;margin-bottom:8px;text-align:center;}
    .loading-subtitle{font-size:14px;color:#666;font-weight:500;text-align:center;
        padding:0 20px;line-height:1.5;}
    </style>
    <div class="loading-screen">
        <div class="spinner-large"></div>
        <div class="loading-title">🚀 TOM'STT AI is Working...</div>
        <div class="loading-subtitle">
            AI sedang menyesuaikan format dengan dokumen acuan Anda.<br>
            Mohon jangan tutup atau keluar dari halaman ini.
        </div>
    </div>
    """, unsafe_allow_html=True)

    ai_result      = None
    error_terakhir = ""

    try:
        active_keys = get_active_keys("Gemini")

        if not active_keys:
            st.error("\u274c Semua API Key Gemini sudah habis. Coba lagi besok atau tambahkan key baru di Panel Admin.")
            st.stop()

        for key_data in active_keys:
            try:
                genai.configure(api_key=key_data["key"])
                _tpl_model = key_data.get("model", "gemini-2.5-flash")
                model    = genai.GenerativeModel(_tpl_model)
                # Nonaktifkan thinking mode untuk model yang support thinking_config
                # (gemini-2.5-pro / gemini-2.5-flash, bukan lite).
                # Model lain (gemini-3.x-lite, dll) tidak punya field ini → jangan diisi.
                _THINKING_CAPABLE = ("gemini-2.5-pro", "gemini-2.5-flash")
                _gen_cfg = None
                if any(p in _tpl_model for p in _THINKING_CAPABLE) and "lite" not in _tpl_model:
                    _gen_cfg = {"thinking_config": {"thinking_level": "minimal"}}
                try:
                    response = model.generate_content(
                        f"{prompt_custom}\n\n[TRANSKRIP SUMBER]\n{st.session_state.transcript}",
                        generation_config=_gen_cfg
                    )
                except Exception as _cfg_err:
                    # Fallback: jika model tidak support thinking_config, retry tanpa config
                    if _gen_cfg and "Unknown field" in str(_cfg_err):
                        response = model.generate_content(
                            f"{prompt_custom}\n\n[TRANSKRIP SUMBER]\n{st.session_state.transcript}",
                            generation_config=None
                        )
                    else:
                        raise
                # response.text bisa throw ValueError jika diblokir safety filter
                try:
                    ai_result = response.text
                except (ValueError, AttributeError):
                    try:
                        ai_result = response.candidates[0].content.parts[0].text
                    except Exception:
                        ai_result = None

                if ai_result and ai_result.strip():
                    # Paid: weighted | Free: flat +1
                    _bobot_tmpl = max(1, len(st.session_state.transcript) // 500) if key_data.get("is_paid") else 1
                    increment_api_usage(key_data["id"], key_data["used"], count=_bobot_tmpl)
                    break
                else:
                    _nama_kt = key_data.get("name", "?")
                    error_terakhir = f"Key [{_nama_kt}]: response kosong/diblokir safety filter"
                    ai_result = None

            except Exception as e:
                err_str = str(e)
                _nama_kt = key_data.get("name", "?")
                error_terakhir = f"Key [{_nama_kt}]: {type(e).__name__}: {err_str[:150]}"

                # Auto-exhaust key jika 429
                if "429" in err_str or "ResourceExhausted" in err_str or "quota" in err_str.lower():
                    try:
                        db.collection('api_keys').document(key_data["id"]).update({
                            "used": key_data["limit"]
                        })
                    except Exception:
                        pass

                # Toast detail untuk admin
                if st.session_state.user_role == "admin":
                    st.toast(f"⚠️ Key [{_nama_kt}] gagal: {err_str[:80]}", icon="🔑")

                continue

    finally:
        # Loading overlay SELALU dibersihkan apapun yang terjadi
        loading.empty()

    if not ai_result:
        st.error(
            f"\u274c **Gagal memproses AI.** Saldo & Kuota Anda **AMAN** (tidak dipotong ganda).\n\n"
            f"**Detail:** `{error_terakhir or 'Semua key gagal tanpa pesan error'}`"
        )
        st.stop()

    # ── Simpan hasil ──────────────────────────────────────────────
    st.session_state.ai_result = ai_result
    safe_dok_name              = nama_dok.replace(" ", "_").replace("/", "_")
    st.session_state.ai_prefix = f"Custom_{safe_dok_name}_"

    hak_arsip = False
    if is_admin or is_b2b_custom:
        hak_arsip = True
    else:
        sys_conf_arsip = sys_config.get(
            "archive_allowed_packages",
            ["EKSEKUTIF", "VIP", "ENTERPRISE", "AIO 10 JAM", "AIO 30 JAM", "AIO 100 JAM"]
        )
        for pkt in user_info.get("inventori", []):
            if any(ap in pkt["nama"].upper() for ap in sys_conf_arsip):
                hak_arsip = True
                break

    db.collection('users').document(st.session_state.current_user).update({
        "draft_transcript": st.session_state.transcript,
        "draft_filename":   st.session_state.filename,
        "draft_ai_result":  ai_result,
        "draft_ai_prefix":  st.session_state.ai_prefix
    })

    vid_save      = user_info.get("active_corporate_voucher")
    sec_mode_save = "Normal"
    if vid_save:
        v_doc_s       = db.collection('vouchers').document(vid_save).get().to_dict() or {}
        sec_mode_save = v_doc_s.get("security_mode", "Normal")

    if sec_mode_save != "Zero Retention (v0)":
        db.collection('users').document(st.session_state.current_user).collection('history').add({
            "filename":   st.session_state.filename,
            "transcript": st.session_state.transcript,
            "ai_result":  ai_result,
            "ai_prefix":  st.session_state.ai_prefix,
            "hak_arsip":  hak_arsip,
            "created_at": firestore.SERVER_TIMESTAMP
        })

    st.success("✔ **Dokumen berhasil dibuat sesuai format template Anda!**")
    time.sleep(1)
    st.rerun()